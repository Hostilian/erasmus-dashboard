import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Ensure stdout handles encoding
sys.stdout.reconfigure(encoding='utf-8')

def create_letter(filename, university_name, location, body_paragraphs):
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # 1. Header (Sender Info)
    p_sender = doc.add_paragraph()
    p_sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sender = p_sender.add_run(
        "Eren Ozturk\n"
        "Faculty of Economics and Management (PEF)\n"
        "Czech University of Life Sciences Prague (CZU)\n"
        "Prague, Czech Republic\n"
        "Email: hostiliann@gmail.com\n"
        "Date: May 27, 2026"
    )
    run_sender.font.size = Pt(10)
    
    # Space
    doc.add_paragraph()
    
    # 2. Recipient Info
    p_recipient = doc.add_paragraph()
    p_recipient.add_run(
        "To the Erasmus Selection Committee,\n"
        f"Office of International Relations\n"
        f"{university_name}\n"
        f"{location}"
    ).bold = True
    
    # Space
    doc.add_paragraph()
    
    # 3. Subject
    p_subject = doc.add_paragraph()
    run_sub = p_subject.add_run(f"Subject: Motivation Letter for Erasmus+ Exchange Mobility - Fall Semester 2026/2027")
    run_sub.bold = True
    
    # Space
    doc.add_paragraph()
    
    # 4. Body Paragraphs
    for para in body_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(12)
        p.add_run(para)
        
    # Space
    doc.add_paragraph()
    
    # 5. Sign-off
    p_sign = doc.add_paragraph()
    p_sign.add_run(
        "Sincerely,\n\n\n"
        "Eren Ozturk\n"
        "BSc Informatics Student, CZU Prague"
    )
    
    doc.save(filename)
    print(f"Successfully generated: {filename}")

# Define paragraph contents for each university
porto_paras = [
    "I am writing to formally express my strong interest in participating in the Erasmus+ exchange program at Universidade Portucalense (UPT) in Porto, Portugal, for the Fall semester of the 2026/2027 academic year. As a BSc Informatics student at the Czech University of Life Sciences Prague (CZU) who also works professionally as a software developer and data analyst, I am eager to combine my technical experience with UPT's outstanding curriculum in Informatics Engineering and Management.",
    "UPT is my top-choice host institution because its course offerings align perfectly with my degree requirements. I plan to enroll in core subjects that serve as direct equivalents to my CZU courses: Sistemas Operativos, Tecnologias Web, Estatística e Análise de Dados, and Contabilidade Geral. The opportunity to study these in UPT's English-supported academic environment will allow me to maintain my double-focus in software engineering and business informatics without delaying my academic progress. Having worked as a developer, I understand that bridging system-level programming with data-driven analytics is crucial, and UPT's syllabus provides precisely this balance.",
    "Beyond academics, Porto is a fantastic cultural and personal fit. Having lived and worked in Prague as a Turkish national, I am highly independent and experienced in managing international relocations. I will be traveling with my pet cat, which requires diligent logistical and veterinary planning (including managing TAP Air Portugal's cabin weight limits), and I have already secured the necessary medical certifications. In my free time, I look forward to engaging with Porto's vibrant student community, participating in ESN activities, and training at local bouldering facilities like São Rock Climbing and Proa Climbing Center.",
    "I am confident that my technical skills, professional drive, and adaptability make me a strong candidate for this exchange and a positive representative for both CZU Prague and UPT Porto. Thank you for considering my application. I look forward to the possibility of studying at Universidade Portucalense."
]

lleida_paras = [
    "I am writing to formally express my strong interest in participating in the Erasmus+ exchange program at the Universitat de Lleida (UdL), Spain, for the Fall semester of the 2026/2027 academic year. As a BSc Informatics student at the Czech University of Life Sciences Prague (CZU) who is also a working software developer and data analyst, I am excited about the prospect of joining your Escola Politècnica Superior (EPS) and FDET faculties.",
    "UdL is an exceptional match for my academic objectives. I have mapped four courses for my winter semester exchange: Sistemes Operatius, Web dinàmica, Estadística Avançada, and Fonaments de Comptabilitat. These courses provide a direct equivalent to my remaining CZU requirements, combining core systems engineering with advanced statistical analysis and business accounting. Exploring these subjects under the guidance of UdL's faculty will enhance my backend development skills and prepare me for complex system integrations. I look forward to collaborating with local and international students in Lleida.",
    "Lleida represents an ideal environment for my personal growth and active lifestyle. As an international student working in Prague, I have developed strong independence and adaptability. I am fully prepared to manage the relocation logistics, including traveling with my pet cat via Barcelona (BCN) using Vueling's cabin allotment and connecting to Lleida via the high-speed AVE train. Additionally, as a passionate climber, Lleida's proximity to world-class climbing areas like Margalef, Santa Linya, and Oliana, along with local gyms like Boulder Indoor Lleida, makes it a dream destination where I can balance rigorous studies with my athletic training.",
    "I am confident that my academic commitment, professional coding experience, and enthusiasm will make me a valuable addition to the student body at the Universitat de Lleida. Thank you for your time and consideration of my application. I hope to represent CZU Prague as an exchange student at UdL."
]

sofia_paras = [
    "I am writing to formally express my strong interest in participating in the Erasmus+ exchange program at the University of National and World Economy (UNWE) in Sofia, Bulgaria, for the Fall semester of the 2026/2027 academic year. As a BSc Informatics student at the Czech University of Life Sciences Prague (CZU) who also works professionally as a software developer and data analyst, I am eager to combine my technical background with the advanced business informatics curriculum offered at your prestigious institution.",
    "UNWE is my top-choice host university because of its long-standing reputation as a leading center for economic and informatics studies in Southeastern Europe. The Business Informatics program at UNWE aligns perfectly with my academic goals and my remaining CZU curriculum. Specifically, I plan to enroll in courses that serve as equivalents to my core CZU requirements: Operating Systems, Statistical Analysis, and Accounting. Having worked as a developer, I know that masterfully bridging IT systems with economic frameworks is key to building impactful business applications, and the academic package at UNWE provides exactly this combination.",
    "Beyond academics, Sofia represents an excellent cultural and logistical fit. As a non-EU passport holder residing in Prague, I have already navigated international relocation and demonstrated adaptability. I plan to relocate with my cat, which requires responsible planning, and I am fully prepared to manage my accommodation and integration in Bulgaria. Additionally, I look forward to connecting with Sofia's active student and professional community, as well as exploring local recreational climbing and bouldering communities, which is a major personal passion of mine.",
    "I am confident that my technical skills, academic focus, and work ethic make me an excellent representative for both CZU Prague and the Erasmus+ program. Thank you for considering my application. I look forward to the possibility of studying at the University of National and World Economy."
]

sibiu_paras = [
    "I am writing to formally express my strong interest in participating in the Erasmus+ exchange program at the Lucian Blaga University of Sibiu (LBU), Romania, for the Fall semester of the 2026/2027 academic year. As an Informatics student at the Czech University of Life Sciences Prague (CZU) with professional experience as a developer and data analyst, I am excited about the prospect of taking part in your university's English-taught Business Informatics curriculum.",
    "LBU's Business Informatics program is highly appealing to me due to its modern, practical approach to software engineering and systems analysis. I have identified direct matches for my CZU requirements within your course catalog, including UNIX/Computer Networks, Web Technologies, and Mobile Business/ERP. Exploring these subjects in LBU's international classrooms will allow me to deepen my technical expertise in backend systems and database interfaces while earning the 30 ECTS required for my degree progression.",
    "Sibiu is a beautiful historical city known for its welcoming environment and rich multicultural heritage, making it an ideal destination for study and personal growth. As a non-EU student who has successfully lived and worked in Prague, I have proven myself to be highly independent and adaptable. Moving to Sibiu, along with relocating my pet cat, is a step I am fully prepared for. Furthermore, I am eager to join the local student community, participate in university organizations, and explore the nearby climbing facilities and bouldering gym networks in Romania during my free time.",
    "I am confident that my academic commitment, coding background, and adaptability make me a strong candidate for this exchange. Thank you for your time and consideration of my application. I look forward to the opportunity to join the academic community at Lucian Blaga University of Sibiu."
]

kosice_paras = [
    "I am writing to formally express my strong interest in participating in the Erasmus+ exchange program at the Technical University of Košice (TUKE), Slovakia, for the Fall semester of the 2026/2027 academic year. As an Informatics student at the Czech University of Life Sciences Prague (CZU) who works as a software developer and data analyst, I am eager to engage with the advanced technology and engineering-focused programs at TUKE.",
    "TUKE stands out as one of Slovakia's premier technical universities, making it an exceptional fit for my academic path. I am particularly drawn to your English-taught Informatics and Economics courses. At TUKE, I aim to map equivalent modules for UNIX Operating Systems, Web Design, and Statistics, leveraging your university's state-of-the-art laboratory facilities. Additionally, taking business and financial accounting modules at your Faculty of Economics will allow me to maintain my double-focus in IT and business analytics without delaying my graduation.",
    "Slovakia's close cultural ties and geographical proximity to Prague make TUKE an excellent choice for a smooth yet enriching exchange experience. As a non-EU resident in Prague, I have established strong independence. I am fully prepared to manage the logistics of this mobility, including relocating my pet cat. Furthermore, Košice has a vibrant technology scene and a very active outdoor bouldering and sport climbing community, which aligns perfectly with my lifestyle and personal interests.",
    "I believe my combination of professional developer experience and strong academic drive will make me a positive and active addition to your student body. Thank you for considering my application. I hope to represent CZU Prague as an exchange student at the Technical University of Košice."
]

# Generate letters
create_letter("motivation_letter_sofia.docx", "University of National and World Economy", "Sofia, Bulgaria", sofia_paras)
create_letter("motivation_letter_sibiu.docx", "Lucian Blaga University of Sibiu", "Sibiu, Romania", sibiu_paras)
create_letter("motivation_letter_kosice.docx", "Technical University of Košice", "Košice, Slovakia", kosice_paras)
create_letter("motivation_letter_porto.docx", "Universidade Portucalense", "Porto, Portugal", porto_paras)
create_letter("motivation_letter_lleida.docx", "Universitat de Lleida", "Lleida, Spain", lleida_paras)
